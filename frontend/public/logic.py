import re
import io
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

# --- Constants & Helpers ---

KYC_FIELDS = [
    ("name", "Name"),
    ("country", "Country"),
    ("gender", "Gender"),
    ("id_expired", "ID Expired"),
    ("id_expiry", "ID Expiry Date"),
    ("id_type", "ID Type"),
    ("id_number", "ID Number"),
    ("dob", "Date of Birth"),
    ("submit_time", "Submit Time"),
    ("review_time", "Review Time"),
    ("submit_ip", "Submit IP"),
    ("ip_location", "IP Location"),
    ("device_id", "Submit Device ID"),
    ("device_type", "Device Type"),
]

KYC_LABEL_CANONICAL = {label.lower(): label for _key, label in KYC_FIELDS}
KYC_TEMPLATE_LABEL_ALIASES = {
    "身份": "ID Type",
}

def normalize_kyc_template_label(label: str) -> str:
    cleaned = label.strip()
    if not cleaned:
        return cleaned
    alias = KYC_TEMPLATE_LABEL_ALIASES.get(cleaned)
    if not alias:
        alias = KYC_TEMPLATE_LABEL_ALIASES.get(cleaned.lower())
    if alias:
        return alias
    return KYC_LABEL_CANONICAL.get(cleaned.lower(), cleaned)

def _replace_paragraph_text_preserve_format(paragraph, new_text):
    """Replace all text in a paragraph while preserving the first run's formatting (rPr)."""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)

def _replace_cell_text_preserve_format(cell, new_text):
    """Replace text in a table cell while preserving the first paragraph/run's formatting."""
    if cell.paragraphs:
        _replace_paragraph_text_preserve_format(cell.paragraphs[0], new_text)
        # Remove extra paragraphs if any
        for p in cell.paragraphs[1:]:
            p_elem = p._element
            p_elem.getparent().remove(p_elem)
    else:
        cell.text = new_text

def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def insert_paragraph_after(paragraph, text=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para

def document_to_bytes(doc):
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# --- Listing Agreement Logic ---

def replace_placeholders_in_paragraph(paragraph, mapping):
    if not paragraph.runs:
        return

    # Try per-run replacement first (handles most cases, preserves formatting perfectly)
    found_any = False
    for run in paragraph.runs:
        original = run.text
        new_text = original
        for key, value in mapping.items():
            patterns = [
                f"{{{{{key}}}}}",
                f"{{{{ {key} }}}}",
                f"{{{{{key} }}}}",
                f"{{{{ {key}}}}}",
            ]
            for pattern in patterns:
                new_text = new_text.replace(pattern, str(value or ""))
        if new_text != original:
            run.text = new_text
            found_any = True
    # Per-run replacement done. Now check for cross-run placeholders
    # (do NOT return early here — a single paragraph may have both
    #  single-run and cross-run placeholders, e.g. {{company}} in one run
    #  and {{Jurisdiction}} split across multiple runs)

    # Cross-run replacement: placeholder is split across runs
    # Join all run texts and check if any placeholder exists
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return

    new_full = full_text
    for key, value in mapping.items():
        patterns = [
            f"{{{{{key}}}}}",
            f"{{{{ {key} }}}}",
            f"{{{{{key} }}}}",
            f"{{{{ {key}}}}}",
        ]
        for pattern in patterns:
            new_full = new_full.replace(pattern, str(value or ""))

    if new_full == full_text:
        return

    # Distribute replaced text back across runs, preserving each run's formatting (rPr)
    runs = paragraph.runs
    if len(runs) == 1:
        runs[0].text = new_full
        return

    # Build char-to-run mapping for the original full_text
    run_boundaries = []  # list of (start, end, run_index)
    pos = 0
    for i, run in enumerate(runs):
        length = len(run.text)
        run_boundaries.append((pos, pos + length, i))
        pos += length

    # Find all placeholder spans in the original text and their replacements
    replacements = []  # (start, end, replacement_value)
    temp = full_text
    for key, value in mapping.items():
        patterns = [
            f"{{{{{key}}}}}",
            f"{{{{ {key} }}}}",
            f"{{{{{key} }}}}",
            f"{{{{ {key}}}}}",
        ]
        for pattern in patterns:
            idx = 0
            while True:
                found = temp.find(pattern, idx)
                if found == -1:
                    break
                replacements.append((found, found + len(pattern), str(value or "")))
                # Replace in temp so we don't double-match
                temp = temp[:found] + "\x00" * len(pattern) + temp[found + len(pattern):]
                idx = found + len(pattern)

    if not replacements:
        return

    replacements.sort(key=lambda r: r[0])

    # Rebuild text for each run using skip_until to handle cross-run placeholders
    new_run_texts = [''] * len(runs)
    repl_idx = 0
    skip_until = 0  # global: chars before this position are part of an already-handled replacement

    for ri, (r_start, r_end, _) in enumerate(run_boundaries):
        parts = []
        # Start from the later of r_start or skip_until (skipping chars consumed by a previous replacement)
        cur = max(r_start, skip_until)
        while cur < r_end:
            if repl_idx < len(replacements):
                rp_start, rp_end, rp_val = replacements[repl_idx]
                if rp_start >= cur and rp_start < r_end:
                    # Replacement starts within this run
                    if rp_start > cur:
                        parts.append(full_text[cur:rp_start])
                    parts.append(rp_val)
                    repl_idx += 1
                    skip_until = rp_end
                    cur = rp_end
                    if cur >= r_end:
                        break
                elif rp_start >= r_end:
                    # No more replacements in this run
                    parts.append(full_text[cur:r_end])
                    cur = r_end
                else:
                    # rp_start < cur: shouldn't normally happen, advance
                    repl_idx += 1
            else:
                parts.append(full_text[cur:r_end])
                cur = r_end
        new_run_texts[ri] = ''.join(parts)

    # Apply new texts to runs (preserving each run's rPr/formatting)
    for i, run in enumerate(runs):
        run.text = new_run_texts[i]

def replace_placeholders(doc, mapping):
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, mapping)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_placeholders_in_paragraph(paragraph, mapping)
        for paragraph in section.footer.paragraphs:
            replace_placeholders_in_paragraph(paragraph, mapping)
        # Handle tables in headers/footers if necessary

def remove_technical_fee_clause(doc):
    start_regex = re.compile(r"\bA\s+Technical\s+Fee\b", flags=re.IGNORECASE)
    clause_started = False
    paragraphs = list(doc.paragraphs)
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if not clause_started and start_regex.search(text):
            clause_started = True
        if clause_started:
            if re.match(r"^[dD]\.\s", text) or re.match(r"^(IV|IV\.)\b", text):
                clause_started = False
                continue
            remove_paragraph(paragraph)

def wallets_to_text(wallets):
    ordered = [
        ("erc20", "USDT-ERC20"),
        ("bsc", "USDT-BSC"),
        ("trc20", "USDT-TRC20"),
        ("solana", "USDT/USDC-Solana"),
    ]
    lines = []
    for key, label in ordered:
        val = wallets.get(key, "").strip()
        if val:
            lines.append(f"{label}: {val}")
    return "\n".join(lines)

def update_wallet_clause(doc, wallets, wallet_text):
    target_index = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if "A Technical Fee" in paragraph.text:
            target_index = idx
            break
    if target_index is None:
        return

    paragraph = doc.paragraphs[target_index]
    base_text = paragraph.text.replace("\u00A0", " ").strip()
    confirm_regex = re.compile(r"\s*BitMart confirms the wallet addresses as below:\s*", flags=re.IGNORECASE)
    base_text = confirm_regex.sub(" ", base_text).strip()
    if base_text and not base_text.endswith("."):
        base_text += "."
        
    has_wallets = any(wallets.values()) or bool(wallet_text)
    
    if has_wallets:
        new_text = f"{base_text} BitMart confirms the wallet addresses as below:"
    else:
        new_text = base_text

    _replace_paragraph_text_preserve_format(paragraph, new_text)

    # Find end of block and remove old content
    end_index = None
    for j in range(target_index + 1, len(doc.paragraphs)):
        text = doc.paragraphs[j].text.strip()
        if text.startswith("The Exchange will acknowledge"):
            end_index = j
            break
        if re.match(r"^[dD]\.\s", text) or re.match(r"^(IV|IV\.)\b", text):
            end_index = j
            break
    if end_index is None:
        end_index = target_index + 1

    block_start = target_index + 1
    block_end = end_index # exclusive

    # If no wallets, remove everything between
    if not has_wallets:
        for idx in range(block_end - 1, block_start - 1, -1):
            remove_paragraph(doc.paragraphs[idx])
        return

    # If wallets, check for existing placeholder
    final_wallet_text = wallet_text if wallet_text else wallets_to_text(wallets)
    wallet_lines = final_wallet_text.splitlines()
    
    # Try to find existing placeholder or addresses to replace
    placeholder_para = None
    # Checking paragraphs in the block
    # Note: doc.paragraphs list changes as we remove, but here we scan first
    # We should re-scan or be careful with indices. 
    # Actually safe if we don't modify yet.
    
    # But for simplicity, let's just remove the old block and insert new paragraphs
    for idx in range(block_end - 1, block_start - 1, -1):
         remove_paragraph(doc.paragraphs[idx])
         
    # Now insert new lines
    cursor = paragraph # The "Technical Fee" paragraph
    # Insert blank line
    cursor = insert_paragraph_after(cursor, "")
    for line in wallet_lines:
        cursor = insert_paragraph_after(cursor, line)
    insert_paragraph_after(cursor, "") # Trailing blank line

def generate_listing_agreement(template_bytes, data):
    doc = Document(io.BytesIO(bytes(template_bytes)))
    
    # Build mapping: add aliases so form field names map to template placeholder names
    # Template uses: {{Jurisdiction}}, {{date}}, {{listing}}, {{name1}}, {{Wallets}}
    # Form sends:    jurisdiction,     signdate, listingdate, signname,   wallets
    mapping = dict(data)
    alias_map = {
        "Jurisdiction": data.get("jurisdiction", ""),
        "date": data.get("signdate", ""),
        "listing": data.get("listingdate", ""),
        "name1": data.get("signname", ""),
        "Wallets": data.get("walletText", "") or data.get("wallets", ""),
    }
    for key, value in alias_map.items():
        if key not in mapping:
            mapping[key] = value
    
    # Placeholders
    replace_placeholders(doc, mapping)
    
    # Technical Fee
    if not data.get("includeTechnicalFee", True):
        remove_technical_fee_clause(doc)

    # Wallets
    update_wallet_clause(doc, data.get("wallets", {}), data.get("walletText", ""))
    
    return document_to_bytes(doc)

# --- KYC Logic ---

KYC_LABELS = {
    "account_id": ["account id", "account id with bitmart", "cid"],
    "name": ["\u59d3\u540d", "name", "full name"],
    "country": ["\u56fd\u5bb6", "\u56fd\u7c4d", "country"],
    "gender": ["\u6027\u522b", "gender"],
    "id_expired": [
        "\u8bc1\u4ef6\u662f\u5426\u8fc7\u671f",
        "\u662f\u5426\u8fc7\u671f",
        "\u8bc1\u4ef6\u8fc7\u671f",
        "\u8bc1\u4ef6\u662f\u5426\u6709\u6548",
        "id expired",
        "expired",
    ],
    "id_expiry": [
        "\u8bc1\u4ef6\u8fc7\u671f\u65f6\u95f4",
        "\u8fc7\u671f\u65f6\u95f4",
        "\u8bc1\u4ef6\u5230\u671f\u65f6\u95f4",
        "\u5230\u671f\u65f6\u95f4",
        "\u6709\u6548\u671f",
        "\u8bc1\u4ef6\u6709\u6548\u671f",
        "id expiry date",
        "expiry date",
        "expiration date",
    ],
    "id_type": [
        "\u8bc1\u4ef6\u7c7b\u578b",
        "\u8bc1\u4ef6\u7c7b\u522b",
        "\u8bc1\u4ef6\u79cd\u7c7b",
        "id type",
        "document type",
    ],
    "id_number": [
        "\u8bc1\u4ef6\u53f7",
        "\u8bc1\u4ef6\u53f7\u7801",
        "\u8bc1\u4ef6\u7f16\u53f7",
        "id number",
        "document number",
    ],
    "dob": ["\u751f\u65e5", "\u51fa\u751f\u65e5\u671f", "date of birth", "dob"],
    "submit_time": ["\u63d0\u4ea4\u65f6\u95f4", "\u63d0\u4ea4\u65e5\u671f", "submit time"],
    "review_time": ["\u5ba1\u6838\u65f6\u95f4", "\u5ba1\u6838\u65e5\u671f", "review time"],
    "submit_ip": [
        "\u63d0\u4ea4IP",
        "\u63d0\u4ea4 IP",
        "\u63d0\u4ea4ip",
        "\u63d0\u4ea4IP\u5730\u5740",
        "submit ip",
    ],
    "ip_location": [
        "IP\u5f52\u5c5e\u5730",
        "IP \u5f52\u5c5e\u5730",
        "IP\u6240\u5728\u5730",
        "IP\u6240\u5c5e\u5730",
        "ip\u5f52\u5c5e\u5730",
        "ip location",
    ],
    "device_id": [
        "\u63d0\u4ea4\u8bbe\u5907",
        "\u8bbe\u5907ID",
        "\u8bbe\u5907id",
        "\u8bbe\u5907\u7f16\u53f7",
        "\u8bbe\u5907\u6807\u8bc6",
        "device id",
        "submit device",
        "submit device id",
        "device identifier",
    ],
    "device_type": ["\u8bbe\u5907\u7c7b\u578b", "\u8bbe\u5907\u7c7b\u522b", "device type"],
    "channel": [
        "\u8ba4\u8bc1\u6e20\u9053",
        "\u8ba4\u8bc1\u65b9\u5f0f",
        "channel",
        "verification channel",
        "kyc channel",
    ],
}

def clean_line(line):
    line = line.strip()
    if line in {"-", "\u2014", "\u2013"}:
        return line
    line = re.sub(r"^(?:[-\u2013\u2014]\s+|[\u2022\u00b7]\s+)", "", line)
    return line.strip()

def normalize_kyc_value(key, value):
    value = value.strip()
    if not value: return value
    # Simplified normalization - add more if needed
    if key == "device_id":
        value = re.sub(r"^(id|device\s*id)\s*[:\uff1a]\s*", "", value, flags=re.IGNORECASE)
    return value

def is_kyc_label_line(line):
    candidate = line.strip()
    if not candidate: return False
    for labels in KYC_LABELS.values():
        for label in labels:
            if re.match(rf"^{re.escape(label)}\s*[:\uff1a]?\s*$", candidate, flags=re.IGNORECASE):
                return True
    return False

def kyc_label_key(line):
    candidate = line.strip()
    if not candidate: return None
    for key, labels in KYC_LABELS.items():
        for label in labels:
             if re.match(rf"^{re.escape(label)}\s*[:\uff1a]?\s*$", candidate, flags=re.IGNORECASE):
                 return key
    return None

def extract_kyc_inline_pairs(line, result):
    occurrences = []
    for key, labels in KYC_LABELS.items():
        for label in labels:
            pattern = rf"{re.escape(label)}(?=\s*[:\uff1a]|\s+|$)\s*[:\uff1a]?"
            for match in re.finditer(pattern, line, flags=re.IGNORECASE):
                occurrences.append((match.start(), match.end(), key))
    if not occurrences:
        return
    occurrences.sort(key=lambda item: item[0])
    for idx, (_start, end, key) in enumerate(occurrences):
        next_start = occurrences[idx + 1][0] if idx + 1 < len(occurrences) else len(line)
        raw_value = line[end:next_start].strip()
        raw_value = raw_value.lstrip(":：").strip()
        # Clean trailing labels if they were part of the next match (regex lookahead isn't perfect here)
        # Actually splitting by occurrences is safer.
        if not raw_value or is_kyc_label_line(raw_value):
            continue
        value = normalize_kyc_value(key, raw_value)
        if key not in result or len(value) > len(result[key]):
            result[key] = value

def extract_kyc_fields(text):
    lines = [clean_line(line) for line in text.splitlines()]
    lines = [line for line in lines if line]
    result = {}
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        extract_kyc_inline_pairs(line, result)
        matched = False
        for key, labels in KYC_LABELS.items():
             for label in labels:
                 pattern = rf"^{re.escape(label)}(?=\s*[:\uff1a]|\s+|$)\s*[:\uff1a]?\s*(.*)$"
                 match = re.match(pattern, line, flags=re.IGNORECASE)
                 if not match: continue
                 
                 value = match.group(1).strip()
                 if value and is_kyc_label_line(value): value = ""
                 
                 if not value:
                     j = idx + 1
                     while j < len(lines) and not lines[j].strip(): j += 1
                     while j < len(lines):
                         label_key = kyc_label_key(lines[j])
                         if label_key is None: break
                         if label_key != key: break
                         j += 1
                         while j < len(lines) and not lines[j].strip(): j += 1
                     
                     if j < len(lines) and not is_kyc_label_line(lines[j]):
                         value = lines[j].strip()
                 
                 if value:
                     result[key] = normalize_kyc_value(key, value)
                 matched = True
                 break
             if matched: break
        idx += 1
    return result

def remove_table_row(row):
    row._tr.getparent().remove(row._tr)

def append_kyc_images(doc, images_bytes_list):
    if not images_bytes_list:
        return
    title = doc.add_paragraph("KYC Pictures")
    if title.runs:
        title.runs[0].bold = True
    else:
        title.add_run("KYC Pictures").bold = True
        
    title.paragraph_format.keep_with_next = True
    
    for img_bytes in images_bytes_list:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.keep_together = True
        run = paragraph.add_run()
        run.add_picture(io.BytesIO(bytes(img_bytes)), width=Inches(3.4))

def fill_kyc_document_logic(template_bytes, data, account_id, images_bytes_list):
    doc = Document(io.BytesIO(bytes(template_bytes)))
    
    # Account ID
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().lower().startswith("account id with bitmart"):
            if account_id:
                _replace_paragraph_text_preserve_format(paragraph, f"Account ID with BitMart: {account_id}")
            break

    label_to_key = {label: key for key, label in KYC_FIELDS}
    
    for table in doc.tables:
        for row in list(table.rows):
            if not row.cells: continue
            raw_label = row.cells[0].text.strip()
            label = normalize_kyc_template_label(raw_label)
            lower_label = label.strip().lower()
            
            if lower_label.startswith("kyc picture"):
                remove_table_row(row)
                continue
            if lower_label == "verification channel" or "\u8ba4\u8bc1\u6e20\u9053" in raw_label or "\u8ba4\u8bc1\u65b9\u5f0f" in raw_label:
                remove_table_row(row)
                continue
                
            if label != raw_label:
                _replace_cell_text_preserve_format(row.cells[0], label)
                
            if label in label_to_key:
                key = label_to_key[label]
                _replace_cell_text_preserve_format(row.cells[1], str(data.get(key, "")))

    append_kyc_images(doc, images_bytes_list)
    return document_to_bytes(doc)

