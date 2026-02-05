import os
import re
from datetime import datetime
from io import BytesIO

import streamlit as st

try:
    from docx import Document
    from docx.shared import Inches
except Exception:  # pragma: no cover - surfaced to user in UI
    st.error("Missing dependency: python-docx. Please install requirements first.")
    st.stop()

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

TEMPLATE_FILES = {
    "Company": "BitMart - Token Listing Agreement - TEMPLATE - Company.docx",
    "Company Waive": "BitMart - Token Listing Agreement - TEMPLATE - Company Waive.docx",
}

KYC_TEMPLATE_FILE = "BitMart User - CID - CONFIDENTIAL.docx"
KYC_TEMPLATE_ORDERED = "BitMart User - CID - CONFIDENTIAL - Ordered.docx"

KYC_FIELDS = [
    ("name", "Name"),
    ("country", "Country"),
    ("gender", "Gender"),
    ("id_expired", "ID Expired"),
    ("id_expiry", "ID Expiry Date"),
    ("id_type", "ID Type"),
    ("id_number", "ID Number"),
    ("dob", "Date of Birth"),
    ("submit_time", "Submit Time"),
    ("review_time", "Review Time"),
    ("submit_ip", "Submit IP"),
    ("ip_location", "IP Location"),
    ("device_id", "Submit Device ID"),
    ("device_type", "Device Type"),
]

KYC_LABEL_CANONICAL = {label.lower(): label for _key, label in KYC_FIELDS}
KYC_TEMPLATE_LABEL_ALIASES = {
    "身份": "ID Type",
}

def normalize_kyc_template_label(label: str) -> str:
    cleaned = label.strip()
    if not cleaned:
        return cleaned
    alias = KYC_TEMPLATE_LABEL_ALIASES.get(cleaned)
    if not alias:
        alias = KYC_TEMPLATE_LABEL_ALIASES.get(cleaned.lower())
    if alias:
        return alias
    return KYC_LABEL_CANONICAL.get(cleaned.lower(), cleaned)

KYC_LABELS = {
    "account_id": ["account id", "account id with bitmart", "cid"],
    "name": ["\u59d3\u540d", "name", "full name"],
    "country": ["\u56fd\u5bb6", "\u56fd\u7c4d", "country"],
    "gender": ["\u6027\u522b", "gender"],
    "id_expired": [
        "\u8bc1\u4ef6\u662f\u5426\u8fc7\u671f",
        "\u662f\u5426\u8fc7\u671f",
        "\u8bc1\u4ef6\u8fc7\u671f",
        "\u8bc1\u4ef6\u662f\u5426\u6709\u6548",
        "id expired",
        "expired",
    ],
    "id_expiry": [
        "\u8bc1\u4ef6\u8fc7\u671f\u65f6\u95f4",
        "\u8fc7\u671f\u65f6\u95f4",
        "\u8bc1\u4ef6\u5230\u671f\u65f6\u95f4",
        "\u5230\u671f\u65f6\u95f4",
        "\u6709\u6548\u671f",
        "\u8bc1\u4ef6\u6709\u6548\u671f",
        "id expiry date",
        "expiry date",
        "expiration date",
    ],
    "id_type": [
        "\u8bc1\u4ef6\u7c7b\u578b",
        "\u8bc1\u4ef6\u7c7b\u522b",
        "\u8bc1\u4ef6\u79cd\u7c7b",
        "id type",
        "document type",
    ],
    "id_number": [
        "\u8bc1\u4ef6\u53f7",
        "\u8bc1\u4ef6\u53f7\u7801",
        "\u8bc1\u4ef6\u7f16\u53f7",
        "id number",
        "document number",
    ],
    "dob": ["\u751f\u65e5", "\u51fa\u751f\u65e5\u671f", "date of birth", "dob"],
    "submit_time": ["\u63d0\u4ea4\u65f6\u95f4", "\u63d0\u4ea4\u65e5\u671f", "submit time"],
    "review_time": ["\u5ba1\u6838\u65f6\u95f4", "\u5ba1\u6838\u65e5\u671f", "review time"],
    "submit_ip": [
        "\u63d0\u4ea4IP",
        "\u63d0\u4ea4 IP",
        "\u63d0\u4ea4ip",
        "\u63d0\u4ea4IP\u5730\u5740",
        "submit ip",
    ],
    "ip_location": [
        "IP\u5f52\u5c5e\u5730",
        "IP \u5f52\u5c5e\u5730",
        "IP\u6240\u5728\u5730",
        "IP\u6240\u5c5e\u5730",
        "ip\u5f52\u5c5e\u5730",
        "ip location",
    ],
    "device_id": [
        "\u63d0\u4ea4\u8bbe\u5907",
        "\u8bbe\u5907ID",
        "\u8bbe\u5907id",
        "\u8bbe\u5907\u7f16\u53f7",
        "\u8bbe\u5907\u6807\u8bc6",
        "device id",
        "submit device",
        "submit device id",
        "device identifier",
    ],
    "device_type": ["\u8bbe\u5907\u7c7b\u578b", "\u8bbe\u5907\u7c7b\u522b", "device type"],
    "channel": [
        "\u8ba4\u8bc1\u6e20\u9053",
        "\u8ba4\u8bc1\u65b9\u5f0f",
        "channel",
        "verification channel",
        "kyc channel",
    ],
}

FIELDS = [
    ("company", "company 名称"),
    ("Jurisdiction", "Jurisdiction 管辖地/国"),
    ("address", "address 公司地址"),
    ("date", "合同签署date"),
    ("listing", "listing 上市日"),
    ("token", "token 代币名称"),
    ("amount", "number 数字"),
    ("amountInWords", "money 上市费英文大写"),
    ("name1", "签署人 1"),
]

FIELD_PATTERNS = {
    "company": [
        r"company\s*名称\s*[:：]\s*(.+)",
        r"company\s*name\s*[:：]\s*(.+)",
        r"compny\s*name\s*[:：]\s*(.+)",
        r"companyname\s*[:：]\s*(.+)",
        r"company\s*nm\s*[:：]\s*(.+)",
        r"company\s*[:：]\s*(.+)",
        r"公司名称\s*[:：]\s*(.+)",
    ],
    "Jurisdiction": [
        r"Jurisdiction\s*管辖地/国\s*[:：]\s*(.+)",
        r"Jurisdiction\s*[:：]\s*(.+)",
        r"管辖地/国\s*[:：]\s*(.+)",
    ],
    "address": [
        r"registered\s*address\s*[:：]\s*(.+)",
        r"company\s*address\s*[:：]\s*(.+)",
        r"addr\s*[:：]\s*(.+)",
        r"address\s*公司地址\s*[:：]\s*(.+)",
        r"address\s*[:：]\s*(.+)",
        r"注册地址\s*[:：]\s*(.+)",
        r"公司地址\s*[:：]\s*(.+)",
    ],
    "date": [
        r"合同签署date\s*[:：]\s*(.+)",
        r"agreement\s*date\s*[:：]\s*(.+)",
        r"signing\s*date\s*[:：]\s*(.+)",
        r"date\s*[:：]\s*(.+)",
    ],
    "listing": [
        r"上市日\s*[:：]\s*(.+)",
        r"listing\s*date\s*[:：]\s*(.+)",
        r"latest\s*listing\s*date\s*[:：]\s*(.+)",
        r"listing\s*day\s*[:：]\s*(.+)",
        r"listing\s*time\s*[:：]\s*(.+)",
        r"list\s*date\s*[:：]\s*(.+)",
        r"time\s*to\s*market\s*[:：]\s*(.+)",
        r"最晚上线日期\s*[:：]\s*(.+)",
        r"上线日期\s*[:：]\s*(.+)",
        r"listing\s*[:：]\s*(.+)",
    ],
    "token": [
        r"token\s*代币名称\s*[:：]\s*(.+)",
        r"token\s*ticker\s*[:：]\s*(.+)",
        r"ticker\s*[:：]\s*(.+)",
        r"token\s*symbol\s*[:：]\s*(.+)",
        r"symbol\s*[:：]\s*(.+)",
        r"token\s*name\s*[:：]\s*(.+)",
        r"token\s*[:：]\s*(.+)",
        r"代币名称\s*[:：]\s*(.+)",
    ],
    "amount": [
        r"number\s*数字\s*[:：]\s*(.+)",
        r"amount\s*[:：]\s*(.+)",
        r"listing\s*fee\s*[:：]\s*(.+)",
        r"金额\s*[:：]\s*(.+)",
    ],
    "amountInWords": [
        r"money\s*上市费英文大写\s*[:：]\s*(.+)",
        r"amountinwords\s*[:：]\s*(.+)",
        r"英文大写\s*[:：]\s*(.+)",
    ],
    "name1": [
        r"签署人\s*1\s*[:：]\s*(.+)",
        r"signer\s*name\s*[:：]\s*(.+)",
        r"full\s*legal\s*name\s*[:：]\s*(.+)",
        r"signer\s*[:：]\s*(.+)",
        r"name1\s*[:：]\s*(.+)",
    ],
}

COUNTRY_KEYWORDS = {
    "indonesia": "Indonesia",
    "singapore": "Singapore",
    "hong kong": "Hong Kong",
    "cayman": "Cayman Islands",
    "british virgin islands": "British Virgin Islands",
    "bvi": "British Virgin Islands",
    "seychelles": "Seychelles",
    "united states": "United States",
    "usa": "United States",
    "united kingdom": "United Kingdom",
    "uk": "United Kingdom",
    "st. vincent and the grenadines": "St. Vincent and the Grenadines",
    "st vincent and the grenadines": "St. Vincent and the Grenadines",
}

NUMBER_RE = re.compile(r"(?<!\w)(\d{1,3}(?:,\d{3})+(?:\.\d+)?|\d+(?:\.\d+)?)")
COMPACT_AMOUNT_RE = re.compile(r"(?<!\w)(\d+(?:\.\d+)?)\s*([kKmM])")
ETH_ADDRESS_RE = re.compile(r"\b0x[a-fA-F0-9]{40}\b")
TRON_ADDRESS_RE = re.compile(r"\bT[1-9A-HJ-NP-Za-km-z]{33}\b")
SOL_ADDRESS_RE = re.compile(r"\b[1-9A-HJ-NP-Za-km-z]{32,44}\b")

ONES = [
    "ZERO",
    "ONE",
    "TWO",
    "THREE",
    "FOUR",
    "FIVE",
    "SIX",
    "SEVEN",
    "EIGHT",
    "NINE",
    "TEN",
    "ELEVEN",
    "TWELVE",
    "THIRTEEN",
    "FOURTEEN",
    "FIFTEEN",
    "SIXTEEN",
    "SEVENTEEN",
    "EIGHTEEN",
    "NINETEEN",
]

TENS = [
    "",
    "",
    "TWENTY",
    "THIRTY",
    "FORTY",
    "FIFTY",
    "SIXTY",
    "SEVENTY",
    "EIGHTY",
    "NINETY",
]

SCALES = [
    (1_000_000_000, "BILLION"),
    (1_000_000, "MILLION"),
    (1_000, "THOUSAND"),
]


def int_to_words(n: int) -> str:
    if n == 0:
        return "ZERO"

    def chunk_to_words(num: int) -> list[str]:
        words: list[str] = []
        if num >= 100:
            words.append(ONES[num // 100])
            words.append("HUNDRED")
            num %= 100
        if num >= 20:
            words.append(TENS[num // 10])
            num %= 10
        if 0 < num < 20:
            words.append(ONES[num])
        return words

    words: list[str] = []
    remainder = n
    for scale_value, scale_name in SCALES:
        if remainder >= scale_value:
            scale_chunk = remainder // scale_value
            words.extend(chunk_to_words(scale_chunk))
            words.append(scale_name)
            remainder %= scale_value
    if remainder:
        words.extend(chunk_to_words(remainder))
    return " ".join(words)


def number_str_to_words(num_str: str) -> str:
    raw = num_str.replace(",", "").strip()
    if not raw:
        return ""
    if "." in raw:
        whole, frac = raw.split(".", 1)
        whole_words = int_to_words(int(whole)) if whole else "ZERO"
        frac_words = " ".join(ONES[int(digit)] for digit in frac if digit.isdigit())
        return f"{whole_words} POINT {frac_words}".strip()
    return int_to_words(int(raw))


def format_number_str(num_str: str) -> str:
    raw = num_str.replace(",", "").strip()
    if not raw:
        return ""
    if "." in raw:
        whole, frac = raw.split(".", 1)
    else:
        whole, frac = raw, ""
    try:
        whole_int = int(whole)
    except ValueError:
        return num_str.strip()
    whole_fmt = f"{whole_int:,}"
    return f"{whole_fmt}.{frac}" if frac else whole_fmt


def parse_compact_amount(text: str) -> str:
    match = COMPACT_AMOUNT_RE.search(text)
    if not match:
        return ""
    value = float(match.group(1))
    suffix = match.group(2).lower()
    multiplier = 1000 if suffix == "k" else 1_000_000
    amount = value * multiplier
    if amount.is_integer():
        return str(int(amount))
    return f"{amount:.2f}".rstrip("0").rstrip(".")


def clean_line(line: str) -> str:
    line = line.strip()
    if line in {"-", "\u2014", "\u2013"}:
        return line
    line = re.sub(r"^(?:[-\u2013\u2014]\s+|[\u2022\u00b7]\s+)", "", line)
    return line.strip()


def find_value(patterns: list[str], lines: list[str]) -> str:
    for line in lines:
        for pattern in patterns:
            match = re.search(pattern, line, flags=re.IGNORECASE)
            if match:
                return match.group(1).strip()
    return ""


def find_number_in_line(line: str) -> str:
    matches = NUMBER_RE.findall(line)
    if not matches:
        return ""
    return max(matches, key=len)


def find_amount_candidate(line: str) -> str:
    compact = parse_compact_amount(line)
    if compact:
        return compact
    return find_number_in_line(line)


def extract_amount_and_words(lines: list[str]) -> tuple[str, str]:
    amount = ""
    words = ""
    for line in lines:
        upper = line.upper()
        if re.search(r"number\s*数字|amount\s*[:：]|金额\s*[:：]|listing\s*fee", line, flags=re.IGNORECASE):
            candidate = find_amount_candidate(line)
            if candidate:
                amount = candidate
        if re.search(r"money|英文大写|amountinwords", line, flags=re.IGNORECASE):
            _, _, tail = line.partition(":" if ":" in line else "：")
            candidate_words = tail.strip()
            if candidate_words:
                words = candidate_words
        if ("USDT" in upper or "USD" in upper or "$" in line or "FEE" in upper) and not amount:
            candidate = find_amount_candidate(line)
            if candidate:
                amount = candidate
        if ("USDT" in upper or "USD" in upper or "FEE" in upper) and not words:
            match = re.search(r"\(\s*([A-Z][A-Z\s\-]+)\s*\)", line)
            if match:
                words = match.group(1).strip()
    return amount, words


def detect_technical_fee(text: str) -> bool:
    lower = text.lower()
    if "technical fee" in lower and ("waive" in lower or "no " in lower or "without" in lower):
        return False
    if "no technical fee" in lower or "without technical fee" in lower:
        return False
    if "technical fee" in lower:
        return True
    return True


def format_date(dt: datetime) -> str:
    return f"{dt.strftime('%B')} {dt.day}, {dt.year}"


def parse_date_string(value: str) -> str:
    value = value.strip()
    if not value:
        return value
    # Extract date portion from phrases like "Before Oct 30" / "by Oct 30".
    match = re.search(r"\b(before|by|on or before)\b\s*(.+)$", value, flags=re.IGNORECASE)
    if match:
        value = match.group(2).strip()
    value = re.sub(r"\b([A-Za-z]{3})\.\b", r"\1", value)
    value = re.sub(r"(\d{1,2})(st|nd|rd|th)\b", r"\1", value, flags=re.IGNORECASE)
    value = re.sub(r"\s+", " ", value).strip()
    # If the date has no year (e.g., "Oct 30"), append current year.
    if not re.search(r"\b\d{4}\b", value):
        value = f"{value} {datetime.today().year}"
    for fmt in (
        "%Y-%m-%d",
        "%Y/%m/%d",
        "%Y.%m.%d",
        "%d-%m-%Y",
        "%m/%d/%y",
        "%m/%d/%Y",
        "%d/%m/%Y",
        "%d/%m/%y",
        "%y-%m-%d",
        "%B %d, %Y",
        "%b %d, %Y",
        "%b %d %Y",
        "%B %d %Y",
        "%d %b %Y",
        "%d %B %Y",
        "%d %b, %Y",
        "%d %B, %Y",
        "%b-%d-%Y",
        "%B-%d-%Y",
    ):
        try:
            parsed = datetime.strptime(value, fmt)
            return format_date(parsed)
        except ValueError:
            continue
    return value


def extract_registered_address(lines: list[str]) -> str:
    for idx, line in enumerate(lines):
        match = re.match(
            r"(registered\s*address|company\s*address|address)\s*[:\uFF1A]\s*(.*)",
            line,
            flags=re.IGNORECASE,
        )
        if not match:
            continue
        addr_parts = [match.group(2).strip()] if match.group(2).strip() else []
        j = idx + 1
        while j < len(lines):
            next_line = lines[j].strip()
            if not next_line:
                j += 1
                continue
            if re.match(
                r"(company\s*name|full\s*legal\s*name|signer\s*name|signer|position|email|deflation|"
                r"open\s*trading\s*pair|mm\s*demand|time\s*to\s*market|listing\s*fee|payment\s*address|"
                r"jurisdiction|token)",
                next_line,
                flags=re.IGNORECASE,
            ):
                break
            if re.match(r".+\s*[:\uFF1A].+", next_line):
                break
            addr_parts.append(next_line)
            j += 1
        return " ".join([part for part in addr_parts if part]).strip()
    return ""


def extract_address_after_company(lines: list[str]) -> str:
    for idx, line in enumerate(lines):
        if re.match(r"(company\s*name|company)\s*[:\uFF1A]\s*.+", line, flags=re.IGNORECASE):
            addr_parts: list[str] = []
            j = idx + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if not next_line:
                    j += 1
                    continue
                if re.match(r".+\s*[:\uFF1A].+", next_line):
                    break
                if re.match(
                    r"(token|ticker|symbol|pair|trading\s*pair|listing|date|signer|email|certificate|payment|"
                    r"special|mm\s*demand|jurisdiction|address|registered)\b",
                    next_line,
                    flags=re.IGNORECASE,
                ):
                    break
                addr_parts.append(next_line)
                j += 1
            if addr_parts:
                return " ".join(addr_parts).strip()
    return ""


def clean_signer_name(value: str) -> str:
    value = value.strip()
    if not value:
        return ""
    value = re.sub(r"\s+", " ", value)
    # Remove common role suffixes.
    roles = [
        "MANAGING DIRECTOR",
        "DIRECTOR",
        "CEO",
        "COO",
        "CFO",
        "CTO",
        "PRESIDENT",
        "MANAGER",
    ]
    upper = value.upper()
    for role in roles:
        if upper.endswith(" " + role):
            return value[: -len(role)].strip()
    return value


def extract_signer_name(lines: list[str]) -> str:
    for idx, line in enumerate(lines):
        match = re.match(r"(signer\s*name|full\s*legal\s*name|signer)\s*[:\uFF1A]\s*(.*)", line, flags=re.IGNORECASE)
        if match:
            raw = match.group(2).strip()
            if raw:
                return clean_signer_name(raw)
            # If the label line has no content, read the next non-empty line.
            j = idx + 1
            while j < len(lines):
                candidate = lines[j].strip()
                if candidate:
                    return clean_signer_name(candidate)
                j += 1
    return ""

def infer_jurisdiction(address: str, lines: list[str]) -> str:
    if address:
        lower_addr = address.lower()
        for key, value in COUNTRY_KEYWORDS.items():
            if re.search(rf"\b{re.escape(key)}\b", lower_addr):
                return value
    for line in lines:
        lower_line = line.lower()
        for key, value in COUNTRY_KEYWORDS.items():
            if re.search(rf"\b{re.escape(key)}\b", lower_line):
                return value
    return ""


def extract_fields(text: str) -> dict:
    raw_lines = [clean_line(line) for line in text.splitlines()]
    lines = [line for line in raw_lines if line]
    result: dict[str, str] = {}

    for key, patterns in FIELD_PATTERNS.items():
        value = find_value(patterns, lines)
        if value:
            result[key] = value

    if not result.get("token"):
        pair_match = None
        for line in lines:
            pair_match = re.search(
                r"(pair|trading\s*pair|open\s*trading\s*pair|币对)\s*[:：]?\s*([A-Z0-9]{2,12})\s*/\s*([A-Z0-9]{2,12})",
                line,
                flags=re.IGNORECASE,
            )
            if pair_match:
                result["token"] = pair_match.group(2).upper()
                break

    address = extract_registered_address(lines)
    if address:
        result["address"] = address
    elif "address" not in result:
        fallback_address = extract_address_after_company(lines)
        if fallback_address:
            result["address"] = fallback_address

    signer = extract_signer_name(lines)
    if signer and not result.get("name1"):
        result["name1"] = signer
    elif result.get("name1"):
        result["name1"] = clean_signer_name(result["name1"])

    amount, words = extract_amount_and_words(lines)
    if amount:
        result["amount"] = amount
    if words:
        result["amountInWords"] = words

    if "date" in result:
        result["date"] = parse_date_string(result["date"])
    if "listing" in result:
        result["listing"] = parse_date_string(result["listing"])

    if "Jurisdiction" not in result or not result["Jurisdiction"]:
        inferred = infer_jurisdiction(result.get("address", ""), lines)
        if inferred:
            result["Jurisdiction"] = inferred

    return result


def extract_wallets(text: str) -> dict[str, str]:
    wallets = {"erc20": "", "bsc": "", "trc20": "", "solana": ""}
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    for line in lines:
        lower = line.lower()
        eth_matches = ETH_ADDRESS_RE.findall(line)
        tron_matches = TRON_ADDRESS_RE.findall(line)
        sol_matches = [addr for addr in SOL_ADDRESS_RE.findall(line) if addr not in tron_matches]

        if ("erc20" in lower or "eth" in lower) and eth_matches and not wallets["erc20"]:
            wallets["erc20"] = eth_matches[0]
        if ("bsc" in lower or "bep20" in lower) and eth_matches and not wallets["bsc"]:
            wallets["bsc"] = eth_matches[0]
        if "trc20" in lower and tron_matches and not wallets["trc20"]:
            wallets["trc20"] = tron_matches[0]
        if "sol" in lower and sol_matches and not wallets["solana"]:
            wallets["solana"] = sol_matches[0]

    if not wallets["trc20"]:
        tron_matches = TRON_ADDRESS_RE.findall(text)
        if tron_matches:
            wallets["trc20"] = tron_matches[0]

    eth_matches = ETH_ADDRESS_RE.findall(text)
    for eth in eth_matches:
        if not wallets["erc20"]:
            wallets["erc20"] = eth
            continue
        if not wallets["bsc"] and eth != wallets["erc20"]:
            wallets["bsc"] = eth
            continue

    sol_candidates = [addr for addr in SOL_ADDRESS_RE.findall(text) if addr not in TRON_ADDRESS_RE.findall(text)]
    if not wallets["solana"] and sol_candidates:
        wallets["solana"] = sol_candidates[0]

    return wallets


def wallets_to_text(wallets: dict[str, str]) -> str:
    ordered = [
        ("erc20", "USDT-ERC20"),
        ("bsc", "USDT-BSC"),
        ("trc20", "USDT-TRC20"),
        ("solana", "USDT/USDC-Solana"),
    ]
    lines = []
    for key, label in ordered:
        address = wallets.get(key, "").strip()
        if address:
            lines.append(f"{label}: {address}")
    return "\n".join(lines)


def wallets_from_input(wallet_text: str, fallback_text: str) -> dict[str, str]:
    wallets = extract_wallets(wallet_text) if wallet_text.strip() else {}
    if not any(wallets.values()):
        wallets = extract_wallets(fallback_text)
    return wallets


def normalize_kyc_value(key: str, value: str) -> str:
    value = value.strip()
    if not value:
        return value
    if key == "gender":
        lower = value.lower()
        if (
            value in ["\u5973", "\u5973\u6027", "female", "f"]
            or lower == "female"
            or "\u5973" in value
        ):
            return "Female"
        if (
            value in ["\u7537", "\u7537\u6027", "male", "m"]
            or lower == "male"
            or "\u7537" in value
        ):
            return "Male"
    if key == "id_expired":
        lower = value.lower()
        if value in [
            "\u672a\u8fc7\u671f",
            "\u672a\u5230\u671f",
            "\u6709\u6548",
            "\u5426",
        ] or lower in ["no", "false", "not expired", "valid"]:
            return "No"
        if value in [
            "\u5df2\u8fc7\u671f",
            "\u8fc7\u671f",
            "\u65e0\u6548",
            "\u5931\u6548",
            "\u662f",
        ] or lower in ["yes", "true", "expired", "invalid"]:
            return "Yes"
    if key == "id_type":
        lower = value.lower()
        if "护照" in value or "passport" in lower:
            return "Passport"
        if "居民身份证" in value or "身份证" in value or "id card" in lower or "identity card" in lower:
            return "ID Card"
        if (
            "居留" in value
            or "resident identity" in lower
            or "residence permit" in lower
            or "residence card" in lower
        ):
            return "Resident Identity Card"
        if "驾照" in value or "驾驶证" in value or "driver license" in lower or "driver's license" in lower:
            return "Driver License"
    if key in ["id_expiry", "dob"]:
        formatted = parse_date_string(value)
        if not formatted:
            match = re.search(r"(\\d{4}[/-]\\d{1,2}[/-]\\d{1,2})", value)
            if match:
                formatted = parse_date_string(match.group(1))
        return formatted if formatted else value
    if key == "device_id":
        value = re.sub(r"^(id|device\s*id)\s*[:\uFF1A]\s*", "", value, flags=re.IGNORECASE)
    return value


# Detect whether a line is just a KYC label (no value).
def is_kyc_label_line(line: str) -> bool:
    candidate = line.strip()
    if not candidate:
        return False
    for labels in KYC_LABELS.values():
        for label in labels:
            pattern = rf"^{re.escape(label)}\s*[:\uFF1A]?\s*$"
            if re.match(pattern, candidate, flags=re.IGNORECASE):
                return True
    return False


def kyc_label_key(line: str) -> str | None:
    candidate = line.strip()
    if not candidate:
        return None
    for key, labels in KYC_LABELS.items():
        for label in labels:
            pattern = rf"^{re.escape(label)}\s*[:\uFF1A]?\s*$"
            if re.match(pattern, candidate, flags=re.IGNORECASE):
                return key
    return None


def extract_kyc_inline_pairs(line: str, result: dict[str, str]) -> None:
    occurrences: list[tuple[int, int, str]] = []
    for key, labels in KYC_LABELS.items():
        for label in labels:
            pattern = rf"{re.escape(label)}(?=\s*[:\uFF1A]|\s+|$)\s*[:\uFF1A]?"
            for match in re.finditer(pattern, line, flags=re.IGNORECASE):
                occurrences.append((match.start(), match.end(), key))
    if not occurrences:
        return
    occurrences.sort(key=lambda item: item[0])
    for idx, (_start, end, key) in enumerate(occurrences):
        next_start = occurrences[idx + 1][0] if idx + 1 < len(occurrences) else len(line)
        raw_value = line[end:next_start].strip()
        raw_value = raw_value.lstrip(":：").strip()
        if not raw_value or is_kyc_label_line(raw_value):
            continue
        value = normalize_kyc_value(key, raw_value)
        if key not in result or len(value) > len(result[key]):
            result[key] = value


# Extract KYC fields from mixed Chinese/English text blocks.
def extract_kyc_fields(text: str) -> dict[str, str]:
    lines = [clean_line(line) for line in text.splitlines()]
    lines = [line for line in lines if line]
    result: dict[str, str] = {}
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        extract_kyc_inline_pairs(line, result)
        matched = False
        for key, labels in KYC_LABELS.items():
            for label in labels:
                pattern = rf"^{re.escape(label)}(?=\s*[:\uFF1A]|\s+|$)\s*[:\uFF1A]?\s*(.*)$"
                match = re.match(pattern, line, flags=re.IGNORECASE)
                if not match:
                    continue
                value = match.group(1).strip()
                if value and is_kyc_label_line(value):
                    value = ""
                if not value:
                    j = idx + 1
                    while j < len(lines) and not lines[j].strip():
                        j += 1
                    while j < len(lines):
                        label_key = kyc_label_key(lines[j])
                        if label_key is None:
                            break
                        if label_key != key:
                            break
                        j += 1
                        while j < len(lines) and not lines[j].strip():
                            j += 1
                    if j < len(lines) and not is_kyc_label_line(lines[j]):
                        value = lines[j].strip()
                if value:
                    result[key] = normalize_kyc_value(key, value)
                matched = True
                break
            if matched:
                break
        idx += 1
    return result


# Ensure an ordered English KYC template exists.
def ensure_kyc_template(path: str) -> None:
    if os.path.exists(path):
        return
    doc = Document()
    doc.add_paragraph("Account ID with BitMart:")
    doc.add_paragraph("1) KYC Info")
    table = doc.add_table(rows=len(KYC_FIELDS), cols=2)
    table.style = "Table Grid"
    for idx, (_key, label) in enumerate(KYC_FIELDS):
        table.cell(idx, 0).text = label
        table.cell(idx, 1).text = ""
    doc.save(path)


# Save a docx document to bytes for download.
def document_to_bytes(doc: Document) -> bytes:
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def remove_table_row(row) -> None:
    row._tr.getparent().remove(row._tr)


def append_kyc_images(doc: Document, images) -> None:
    if not images:
        return
    title = doc.add_paragraph("KYC Pictures")
    if title.runs:
        title.runs[0].bold = True
    title.paragraph_format.keep_with_next = True
    for image in images:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.keep_together = True
        run = paragraph.add_run()
        run.add_picture(BytesIO(image.getvalue()), width=Inches(3.4))


# Fill KYC template with extracted data and images.
def fill_kyc_document(template_path: str, data: dict[str, str], account_id: str, images) -> bytes:
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().lower().startswith("account id with bitmart"):
            if account_id:
                paragraph.text = f"Account ID with BitMart: {account_id}"
            break

    label_to_key = {label: key for key, label in KYC_FIELDS}
    for table in doc.tables:
        for row in list(table.rows):
            raw_label = row.cells[0].text.strip()
            label = normalize_kyc_template_label(raw_label)
            lower_label = label.strip().lower()
            if lower_label.startswith("kyc picture"):
                remove_table_row(row)
                continue
            if lower_label == "verification channel" or "\u8ba4\u8bc1\u6e20\u9053" in raw_label or "\u8ba4\u8bc1\u65b9\u5f0f" in raw_label:
                remove_table_row(row)
                continue
            if label != raw_label:
                row.cells[0].text = label
            if label in label_to_key:
                key = label_to_key[label]
                row.cells[1].text = data.get(key, "")

    append_kyc_images(doc, images)

    return document_to_bytes(doc)


def normalize_amount_fields(amount: str, amount_in_words: str) -> tuple[str, str]:
    amount = amount.strip()
    amount_in_words = amount_in_words.strip()
    if amount:
        compact = parse_compact_amount(amount)
        if compact:
            amount = compact
        amount = format_number_str(amount)
        if not amount_in_words:
            amount_in_words = number_str_to_words(amount)
    if amount_in_words:
        amount_in_words = amount_in_words.upper()
    return amount, amount_in_words


def update_amount_in_words() -> None:
    if (
        st.session_state.get("selected_template") == "Company Waive"
        and not st.session_state.get("include_technical_fee", True)
    ):
        st.session_state["field_amount"] = "N/A"
        st.session_state["field_amountInWords"] = "N/A"
        return
    amount_raw = st.session_state.get("field_amount", "")
    amount, amount_in_words = normalize_amount_fields(amount_raw, "")
    st.session_state["field_amount"] = amount
    st.session_state["field_amountInWords"] = amount_in_words


def clear_kyc_form() -> None:
    st.session_state["kyc_text"] = ""
    st.session_state["kyc_account_id"] = ""
    for key, _label in KYC_FIELDS:
        st.session_state[f"kyc_{key}"] = ""
    st.session_state["kyc_files"] = []
    st.session_state["kyc_uploader_nonce"] = st.session_state.get("kyc_uploader_nonce", 0) + 1
    st.session_state.pop("kyc_download", None)


def replace_placeholders_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    if not paragraph.runs:
        return
    text = "".join(run.text for run in paragraph.runs)
    if not text:
        return
    new_text = text
    for key, value in mapping.items():
        placeholders = [
            f"{{{{{key}}}}}",
            f"{{{{ {key} }}}}",
            f"{{{{{key} }}}}",
            f"{{{{ {key}}}}}",
        ]
        for placeholder in placeholders:
            new_text = new_text.replace(placeholder, value)
    if new_text != text:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def replace_placeholders(doc: Document, mapping: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, mapping)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_placeholders_in_paragraph(paragraph, mapping)
        for paragraph in section.footer.paragraphs:
            replace_placeholders_in_paragraph(paragraph, mapping)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholders_in_paragraph(paragraph, mapping)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholders_in_paragraph(paragraph, mapping)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def remove_technical_fee_clause(doc: Document) -> None:
    start_regex = re.compile(r"\bA\s+Technical\s+Fee\b", flags=re.IGNORECASE)
    clause_started = False
    paragraphs = list(doc.paragraphs)
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if not clause_started and start_regex.search(text):
            clause_started = True
        if clause_started:
            if re.match(r"^[dD]\.\s", text) or re.match(r"^(IV|IV\.)\b", text):
                clause_started = False
                continue
            remove_paragraph(paragraph)


def replace_paragraph_regex(doc: Document, pattern: str, replacement: str) -> None:
    regex = re.compile(pattern, flags=re.IGNORECASE | re.DOTALL)

    def process_paragraph(paragraph) -> None:
        if not paragraph.text:
            return
        raw_text = paragraph.text.replace("\u00A0", " ")
        if not regex.search(raw_text):
            return
        new_text = replacement.strip()
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            process_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            process_paragraph(paragraph)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph(paragraph)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph(paragraph)


def insert_paragraph_after(paragraph, text: str = ""):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def update_wallet_clause(doc: Document, wallets: dict[str, str]) -> None:
    target_index = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if "A Technical Fee" in paragraph.text:
            target_index = idx
            break
    if target_index is None:
        return

    paragraph = doc.paragraphs[target_index]
    base_text = paragraph.text.replace("\u00A0", " ").strip()
    confirm_regex = re.compile(r"\s*BitMart confirms the wallet addresses as below:\s*", flags=re.IGNORECASE)
    base_text = confirm_regex.sub(" ", base_text).strip()
    if base_text.endswith("."):
        base_sentence = base_text
    else:
        base_sentence = base_text.rstrip(".") + "."

    has_wallets = any(wallets.values())
    if has_wallets:
        new_text = f"{base_sentence} BitMart confirms the wallet addresses as below:"
    else:
        new_text = base_sentence

    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)

    end_index = None
    for j in range(target_index + 1, len(doc.paragraphs)):
        text = doc.paragraphs[j].text.strip()
        if text.startswith("The Exchange will acknowledge"):
            end_index = j
            break
        if re.match(r"^[dD]\.\s", text) or re.match(r"^(IV|IV\.)\b", text):
            end_index = j
            break
    if end_index is None:
        end_index = target_index + 1

    block_start = target_index + 1
    block_end = end_index

    if not has_wallets:
        for idx in range(block_end - 1, block_start - 1, -1):
            remove_paragraph(doc.paragraphs[idx])
        # Also remove any stray "BitMart confirms..." paragraph if present.
        for p in list(doc.paragraphs):
            if confirm_regex.search(p.text):
                remove_paragraph(p)
        return

    wallet_lines = wallets_to_text(wallets).splitlines()
    labels = ["USDT-ERC20", "USDT-BSC", "USDT-TRC20", "USDT/USDC-Solana"]
    block_paragraphs = doc.paragraphs[block_start:block_end]
    wallet_block_found = False
    placeholder_para = None
    for p in block_paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if "{{Wallets}}" in text:
            placeholder_para = p
            wallet_block_found = True
            break
        if any(label in text for label in labels) or ETH_ADDRESS_RE.search(text) or TRON_ADDRESS_RE.search(text) or SOL_ADDRESS_RE.search(text):
            placeholder_para = p
            wallet_block_found = True
            break

    if wallet_block_found and placeholder_para:
        new_text = "\n".join(wallet_lines)
        if placeholder_para.runs:
            placeholder_para.runs[0].text = new_text
            for run in placeholder_para.runs[1:]:
                run.text = ""
        else:
            placeholder_para.add_run(new_text)
        for p in list(block_paragraphs):
            if p is placeholder_para:
                continue
            t = p.text.strip()
            if not t:
                continue
            if "{{Wallets}}" in t or any(label in t for label in labels) or ETH_ADDRESS_RE.search(t) or TRON_ADDRESS_RE.search(t) or SOL_ADDRESS_RE.search(t):
                remove_paragraph(p)
        return

    for idx in range(block_end - 1, block_start - 1, -1):
        remove_paragraph(doc.paragraphs[idx])
    blank_para = insert_paragraph_after(paragraph, "")
    cursor = blank_para
    for line in wallet_lines:
        cursor = insert_paragraph_after(cursor, line)
    insert_paragraph_after(cursor, "")


def build_mapping(state: dict) -> dict[str, str]:
    mapping = {}
    for key, _label in FIELDS:
        mapping[key] = state.get(key, "").strip()
    return mapping


def sanitize_filename_component(value: str) -> str:
    value = value.strip()
    if not value:
        return ""
    value = re.sub(r'[<>:"/\\\\|?*]', "-", value)
    value = re.sub(r"\s+", " ", value)
    return value.strip(" .-")


st.set_page_config(page_title="Email to Word", layout="wide")

st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;700&family=Roboto:wght@400;500;700&display=swap');

:root {
    --accent: #111111;
    --accent-2: #1f2937;
    --text: #111827;
    --muted: #6b7280;
    --border: #e5e7eb;
    --input-bg: #f3f4f6;
    color-scheme: light;
}

html, body, .stApp, [data-testid="stApp"], [data-testid="stAppViewContainer"] {
    background: #ffffff !important;
    color-scheme: light !important;
}

input, textarea, select,
div[data-baseweb="input"] input,
div[data-baseweb="select"] input {
    color-scheme: light !important;
}

html, body, [class*="css"]  {
    font-family: 'Inter', 'system-ui', '-apple-system', 'BlinkMacSystemFont', 'Segoe UI', 'Roboto', 'Helvetica Neue', 'Arial', sans-serif;
    color: var(--text);
    font-weight: 500;
    text-rendering: optimizeLegibility;
    -webkit-font-smoothing: antialiased;
    -moz-osx-font-smoothing: grayscale;
}

div[data-testid="stAppViewContainer"] {
    background: #ffffff;
}

header[data-testid="stHeader"] {
    display: none;
}

div[data-testid="stToolbar"],
div[data-testid="stDecoration"],
div[data-testid="stStatusWidget"] {
    display: none;
}

.block-container {
    padding-top: 0.75rem;
    padding-bottom: 2rem;
    max-width: 1536px;
    padding-left: 1rem;
    padding-right: 1rem;
}

h1, h2, h3 {
    font-family: 'Inter', 'system-ui', '-apple-system', 'BlinkMacSystemFont', 'Segoe UI', 'Roboto', 'Helvetica Neue', 'Arial', sans-serif;
    letter-spacing: 0.2px;
}

.app-header {
    display: flex;
    align-items: center;
    gap: 24px;
}

.brand {
    display: flex;
    align-items: center;
    gap: 8px;
    font-weight: 700;
    font-size: 20px;
    color: var(--text);
}

.brand img {
    height: 22px;
    width: auto;
}

#top-nav-anchor + div[data-testid="stRadio"],
#top-nav-anchor ~ div[data-testid="stRadio"],
div.st-key-active_page div[data-testid="stRadio"] {
    width: 100%;
    display: flex;
    justify-content: center;
}

#top-nav-anchor + div[data-testid="stRadio"] > div,
#top-nav-anchor ~ div[data-testid="stRadio"] > div,
div.st-key-active_page div[data-testid="stRadio"] > div {
    background: transparent !important;
    border: none !important;
    padding: 0 !important;
    width: 100%;
    box-shadow: none !important;
}

#top-nav-anchor + div[data-testid="stRadio"] div[role="radiogroup"],
#top-nav-anchor ~ div[data-testid="stRadio"] div[role="radiogroup"],
div.st-key-active_page div[data-testid="stRadio"] div[role="radiogroup"] {
    display: inline-flex !important;
    align-items: center;
    justify-content: center;
    gap: 36px;
    width: fit-content;
    max-width: 100%;
    margin: 0 auto;
    flex-wrap: wrap;
}

#top-nav-anchor + div[data-testid="stRadio"] label,
#top-nav-anchor ~ div[data-testid="stRadio"] label,
div.st-key-active_page div[data-testid="stRadio"] label {
    margin-right: 0;
    flex: 0 0 auto;
}

#top-nav-anchor + div[data-testid="stRadio"] input[type="radio"],
#top-nav-anchor ~ div[data-testid="stRadio"] input[type="radio"],
div.st-key-active_page div[data-testid="stRadio"] input[type="radio"] {
    display: none !important;
}

#top-nav-anchor + div[data-testid="stRadio"] label > div,
#top-nav-anchor ~ div[data-testid="stRadio"] label > div,
div.st-key-active_page div[data-testid="stRadio"] label > div {
    background: transparent;
    border: none;
    color: var(--muted);
    font-weight: 700;
    padding: 2px 0 13px 0;
    border-bottom: 3px solid transparent;
    font-size: 20px;
    line-height: 1.2;
}

#top-nav-anchor + div[data-testid="stRadio"] label > input:checked + div,
#top-nav-anchor ~ div[data-testid="stRadio"] label > input:checked + div,
div.st-key-active_page div[data-testid="stRadio"] label > input:checked + div {
    color: var(--text);
    font-weight: 800;
    border-bottom: 4px solid var(--text);
}

@media (max-width: 1024px) {
    #top-nav-anchor + div[data-testid="stRadio"] div[role="radiogroup"],
    #top-nav-anchor ~ div[data-testid="stRadio"] div[role="radiogroup"],
    div.st-key-active_page div[data-testid="stRadio"] div[role="radiogroup"] {
        gap: 22px;
    }
    #top-nav-anchor + div[data-testid="stRadio"] label > div,
    #top-nav-anchor ~ div[data-testid="stRadio"] label > div,
    div.st-key-active_page div[data-testid="stRadio"] label > div {
        font-size: 17px;
    }
}

@media (max-width: 640px) {
    #top-nav-anchor + div[data-testid="stRadio"] div[role="radiogroup"],
    #top-nav-anchor ~ div[data-testid="stRadio"] div[role="radiogroup"],
    div.st-key-active_page div[data-testid="stRadio"] div[role="radiogroup"] {
        gap: 16px;
    }
    #top-nav-anchor + div[data-testid="stRadio"] label > div,
    #top-nav-anchor ~ div[data-testid="stRadio"] label > div,
    div.st-key-active_page div[data-testid="stRadio"] label > div {
        font-size: 15px;
    }
}

.header-divider {
    height: 1px;
    background: var(--border);
    margin: 0.6rem 0 1.2rem;
}


.step-title {
    font-family: 'Inter', 'system-ui', '-apple-system', 'BlinkMacSystemFont', 'Segoe UI', 'Roboto', 'Helvetica Neue', 'Arial', sans-serif;
    font-size: 18px;
    margin: 6px 0 2px 0;
    color: var(--accent);
}

.step-note {
    color: var(--muted);
    font-size: 13px;
    margin-bottom: 8px;
}

.step-title-note-inline {
    display: flex;
    align-items: baseline;
    gap: 8px;
    flex-wrap: wrap;
    margin: 6px 0 8px 0;
}

.step-title-note-inline .step-title,
.step-title-note-inline .step-note {
    margin: 0;
}

div[data-testid="stTextInput"] input {
    background: #ffffff;
    border: 1px solid #d1d5db;
    border-radius: 8px;
    padding: 0.55rem 0.7rem;
    box-shadow: none;
    color: var(--text);
    font-weight: 500;
}

div[data-testid="stTextInput"] input:focus {
    border-color: #9ca3af;
    box-shadow: 0 0 0 2px rgba(229, 231, 235, 0.6);
}

div[data-testid="stTextArea"] textarea {
    background: var(--input-bg);
    border: 1px solid transparent;
    border-radius: 8px;
    padding: 0.65rem 0.75rem;
    box-shadow: none;
    color: var(--text);
    font-weight: 500;
}

div[data-testid="stTextArea"] textarea:focus {
    background: #ffffff;
    border-color: var(--border);
    box-shadow: 0 0 0 2px rgba(229, 231, 235, 0.5);
}

div[data-testid="stTextInput"] label,
div[data-testid="stTextArea"] label {
    font-size: 12px;
    font-weight: 600;
    color: #374151;
}

div[data-testid="stTextInput"] input::placeholder,
div[data-testid="stTextArea"] textarea::placeholder {
    color: #9ca3af;
    opacity: 1;
}

div[data-testid="stRadio"] > div {
    background: rgba(255, 255, 255, 0.9);
    border: 1px solid rgba(0, 0, 0, 0.14);
    border-radius: 10px;
    padding: 10px 12px;
}


div[data-testid="stButton"] > button {
    background: #111111;
    color: white;
    border: none;
    border-radius: 10px;
    padding: 0.6rem 1.2rem;
    font-weight: 600;
    letter-spacing: 0.2px;
    box-shadow: 0 12px 24px rgba(0, 0, 0, 0.16);
}

div[data-testid="stButton"] > button:hover {
    filter: brightness(0.98);
}

.subtle-divider {
    height: 1px;
    background: linear-gradient(90deg, rgba(15, 118, 110, 0.0), rgba(15, 118, 110, 0.25), rgba(15, 118, 110, 0.0));
    margin: 14px 0;
}

.kyc-title {
    font-size: 20px;
    font-weight: 700;
    margin-bottom: 12px;
}

.kyc-upload div[data-testid="stFileUploader"] {
    border: 2px dashed #d1d5db;
    border-radius: 12px;
    padding: 22px;
    background: #f9fafb;
    text-align: center;
}

.kyc-upload div[data-testid="stFileUploader"] section {
    margin: 0;
}

.kyc-upload div[data-testid="stFileUploader"] button {
    background: #ffffff;
    border: 1px solid var(--border);
    color: var(--accent-2);
    border-radius: 6px;
    padding: 6px 14px;
    font-weight: 600;
}

textarea {
    resize: none;
}
</style>
""",
    unsafe_allow_html=True,
)

if "email_text" not in st.session_state:
    st.session_state.email_text = ""

if "include_technical_fee" not in st.session_state:
    st.session_state.include_technical_fee = True

if "selected_template" not in st.session_state:
    st.session_state.selected_template = "Company"
if "wallets" not in st.session_state:
    st.session_state.wallets = {}
if "wallet_text" not in st.session_state:
    st.session_state.wallet_text = ""

if "kyc_text" not in st.session_state:
    st.session_state.kyc_text = ""
if "kyc_account_id" not in st.session_state:
    st.session_state.kyc_account_id = ""
if "kyc_files" not in st.session_state:
    st.session_state.kyc_files = []
if "kyc_uploader_nonce" not in st.session_state:
    st.session_state.kyc_uploader_nonce = 0
for key, _label in KYC_FIELDS:
    kyc_key = f"kyc_{key}"
    if kyc_key not in st.session_state:
        st.session_state[kyc_key] = ""

ensure_kyc_template(os.path.join(BASE_DIR, KYC_TEMPLATE_ORDERED))

for key, _label in FIELDS:
    session_key = f"field_{key}"
    if session_key not in st.session_state:
        if key == "date":
            st.session_state[session_key] = format_date(datetime.today())
        else:
            st.session_state[session_key] = ""

if "active_page" not in st.session_state:
    st.session_state.active_page = "Listing Agreement"

header_left, header_center, header_right = st.columns([1, 6, 1], gap="small")
with header_left:
    st.markdown(
        """
<div class="brand">
  <img src="https://lh3.googleusercontent.com/aida-public/AB6AXuBCIsWGS4u82vcmjPmBLrtOI3ANk0huPKqPqcIwHyBckN_-9WKP_ILcffQaVFijoO_ZQsMGUsWtf_0CALH-LbV2YG-ZyVwqwSpDKh6abZsDqye96Y4UzrzQpYeyUVnBOoDT4WosfixAmGyXQbJpvNFEdEa_QDzrBTT8KuQRI-bs_xDDcG8XNadNerm8m8FGZ3no7UP4hkdq1iXg6CNfJso2vZTDRKT3PcH2PsbhQt4bw9nF8nM5oZ9_lBmZdS4B-Inr_a2IlRxLzEOm" alt="BitMart Logo"/>
  <span>BitMart</span>
</div>
""",
        unsafe_allow_html=True,
    )
with header_center:
    st.markdown('<div id="top-nav-anchor"></div>', unsafe_allow_html=True)
    st.radio(
        "",
        ["Listing Agreement", "KYC"],
        format_func=lambda x: "BitMart Listing Agreement Generator" if x == "Listing Agreement" else "BitMart KYC Generator",
        horizontal=True,
        key="active_page",
        label_visibility="collapsed",
    )
with header_right:
    st.markdown("&nbsp;", unsafe_allow_html=True)

st.markdown('<div class="header-divider"></div>', unsafe_allow_html=True)

active_page = st.session_state.get("active_page", "Listing Agreement")

if active_page.lower() != "kyc":
    col_left, col_right = st.columns([3, 7], gap="large")
    
    with col_left:
        st.markdown(
            '<div class="step-title-note-inline"><span class="step-title">Step 1 · 邮件输入</span><span class="step-note">-粘贴邮件内容，自动识别并填充字段。</span></div>',
            unsafe_allow_html=True,
        )
        email_text = st.text_area("", height=800, key="email_text", placeholder="在此粘贴邮件内容...")
        if st.button("提取信息 (Analyze)", use_container_width=True):
            parsed = extract_fields(email_text)
            parsed["date"] = format_date(datetime.today())
            if not parsed:
                st.warning("未识别到字段，请手动填写右侧表单。")
            if parsed.get("amount") or parsed.get("amountInWords"):
                normalized_amount, normalized_words = normalize_amount_fields(
                    parsed.get("amount", ""), parsed.get("amountInWords", "")
                )
                if normalized_amount:
                    parsed["amount"] = normalized_amount
                if normalized_words:
                    parsed["amountInWords"] = normalized_words
            for key, value in parsed.items():
                session_key = f"field_{key}"
                if value:
                    st.session_state[session_key] = value
            detected_fee = detect_technical_fee(email_text)
            st.session_state.include_technical_fee = detected_fee
            st.session_state.wallets = extract_wallets(email_text)
            st.session_state.wallet_text = wallets_to_text(st.session_state.wallets)
    
    with col_right:
        st.markdown(
            '<div class="step-title-note-inline"><span class="step-title">Step 2 · 选择模板</span><span class="step-note">-Company 或 Company Waive 模板。</span></div>',
            unsafe_allow_html=True,
        )
        template_choice = st.radio(
            "",
            list(TEMPLATE_FILES.keys()),
            index=list(TEMPLATE_FILES.keys()).index(st.session_state.selected_template),
            horizontal=True,
            label_visibility="collapsed",
        )
        st.session_state.selected_template = template_choice
        if st.session_state.selected_template == "Company Waive":
            st.markdown('<div class="step-note">Company Waive 可选择是否保留 Technical Fee 条款。</div>', unsafe_allow_html=True)
            include_fee = st.radio(
                "Technical Fee 条款",
                ["包含 (Keep)", "不包含 (Remove c.)"],
                index=0 if st.session_state.include_technical_fee else 1,
                horizontal=True,
            )
            st.session_state.include_technical_fee = include_fee.startswith("包含")
            if not st.session_state.include_technical_fee:
                st.session_state["field_amount"] = "N/A"
                st.session_state["field_amountInWords"] = "N/A"
        else:
            st.session_state.include_technical_fee = True
    
        st.markdown('<div class="subtle-divider"></div>', unsafe_allow_html=True)
        st.markdown(
            '<div class="step-title-note-inline"><span class="step-title">Step 3 · 人工核对</span><span class="step-note">-补全或修改必填字段。</span></div>',
            unsafe_allow_html=True,
        )
    
        for row_start in range(0, len(FIELDS), 3):
            row_fields = FIELDS[row_start : row_start + 3]
            row_columns = st.columns(3, gap="large")
            for column, (key, label) in zip(row_columns, row_fields):
                with column:
                    if key == "amount":
                        st.text_input(
                            label,
                            key="field_amount",
                            on_change=update_amount_in_words,
                            disabled=(
                                st.session_state.selected_template == "Company Waive"
                                and not st.session_state.include_technical_fee
                            ),
                        )
                    elif key == "amountInWords":
                        st.text_input(
                            label,
                            key="field_amountInWords",
                            disabled=(
                                st.session_state.selected_template == "Company Waive"
                                and not st.session_state.include_technical_fee
                            ),
                        )
                    else:
                        st.text_input(label, key=f"field_{key}")
    
        st.text_area(
            "钱包地址 Wallets",
            key="wallet_text",
            height=160,
            placeholder="USDT-ERC20: 0x...\nUSDT-BSC: 0x...\nUSDT-TRC20: T...\nUSDT/USDC-Solana: ...",
        )
    
        st.markdown('<div class="subtle-divider"></div>', unsafe_allow_html=True)
        st.markdown('<div class="step-title">Step 4 · 生成文档</div>', unsafe_allow_html=True)
        st.markdown('<div class="step-note">生成后可直接下载 Word 文档。</div>', unsafe_allow_html=True)
    
        if st.button("生成 Word 文档 (Generate)", use_container_width=True):
            state_values = {key: st.session_state.get(f"field_{key}", "") for key, _ in FIELDS}
            if st.session_state.selected_template == "Company Waive" and not st.session_state.include_technical_fee:
                state_values["amount"] = "N/A"
                state_values["amountInWords"] = "N/A"
            amount, amount_in_words = normalize_amount_fields(
                state_values.get("amount", ""), state_values.get("amountInWords", "")
            )
            state_values["amount"] = amount
            state_values["amountInWords"] = amount_in_words
            mapping = build_mapping(state_values)
            wallets = wallets_from_input(
                st.session_state.get("wallet_text", ""),
                st.session_state.get("email_text", ""),
            )
            mapping["Wallets"] = wallets_to_text(wallets) if any(wallets.values()) else ""
    
            template_file = TEMPLATE_FILES[st.session_state.selected_template]
            template_path = os.path.join(BASE_DIR, template_file)
            if not os.path.exists(template_path):
                st.error(f"模板文件不存在: {template_file}")
                st.stop()
    
            doc = Document(template_path)
            replace_placeholders(doc, mapping)
            if not st.session_state.include_technical_fee:
                remove_technical_fee_clause(doc)
                replace_paragraph_regex(
                    doc,
                    r"Developer agrees to coordinate such promotional activities.*?communities"
                    r"(?:,\s*with a budget of .*? USDT\.)?\s*"
                    r"The Developer agrees to make such budget available before the listing date of the Token\.",
                    "Developer agrees to coordinate such promotional activities and to sponsor listing related "
                    "campaigns on the Exchange, including, but not limited to, campaigns on social media platforms, "
                    "press releases, promotional events in communities. The Developer agrees to make such budget "
                    "available before the listing date of the Token.",
                )
            else:
                update_wallet_clause(doc, wallets)
    
            safe_company = sanitize_filename_component(mapping.get("company", "Document")) or "Document"
            safe_token = sanitize_filename_component(mapping.get("token", ""))
            token_prefix = f"{safe_token}-" if safe_token else ""
            filename = f"{token_prefix}{safe_company} - Agreement.docx"
            doc_bytes = document_to_bytes(doc)
            st.session_state["agreement_download"] = {"filename": filename, "bytes": doc_bytes}
            st.success("已生成文档，请下载。")

        download_payload = st.session_state.get("agreement_download")
        if download_payload:
            st.download_button(
                "下载 Word 文档",
                data=download_payload["bytes"],
                file_name=download_payload["filename"],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="agreement_download_btn",
            )
    

else:
    kyc_left, kyc_right = st.columns([3, 7], gap="large")

    with kyc_left:
        st.markdown('<div style="font-size:12px;font-weight:600;color:#374151;">KYC Raw Text</div>', unsafe_allow_html=True)
        kyc_text = st.text_area(
            "",
            height=360,
            key="kyc_text",
            placeholder="Paste KYC raw text here...",
            label_visibility="collapsed",
        )
        st.markdown('<div class="kyc-upload">', unsafe_allow_html=True)
        kyc_uploader_key = f"kyc_files_uploader_{st.session_state.get('kyc_uploader_nonce', 0)}"
        uploaded_images = st.file_uploader(
            "KYC Pictures",
            type=["png", "jpg", "jpeg"],
            accept_multiple_files=True,
            key=kyc_uploader_key,
            label_visibility="collapsed",
        )
        st.session_state["kyc_files"] = uploaded_images or []
        st.markdown("</div>", unsafe_allow_html=True)

        if st.button("Analyze", use_container_width=True, key="kyc_analyze"):
            kyc_parsed = extract_kyc_fields(kyc_text)
            for key, _label in KYC_FIELDS:
                if key in kyc_parsed:
                    st.session_state[f"kyc_{key}"] = kyc_parsed[key]
            if kyc_parsed.get("account_id"):
                st.session_state.kyc_account_id = kyc_parsed["account_id"]
            if not kyc_parsed:
                st.warning("未识别到KYC字段，请检查输入格式。")
            else:
                st.success(f"已识别 {len(kyc_parsed)} 项 KYC 信息。")
        st.button(
            "一键清空",
            use_container_width=True,
            key="kyc_clear",
            on_click=clear_kyc_form,
        )

    with kyc_right:
        st.markdown('<div class="kyc-title">KYC Data Fields</div>', unsafe_allow_html=True)
        if st.session_state.get("kyc_account_id"):
            st.text_input("Account ID with BitMart", key="kyc_account_id")

        if st.session_state.get("kyc_id_type"):
            st.session_state["kyc_id_type"] = normalize_kyc_value(
                "id_type", st.session_state.get("kyc_id_type", "")
            )

        row1 = st.columns(3, gap="large")
        row1[0].text_input("Name", key="kyc_name")
        row1[1].text_input("Country", key="kyc_country")
        row1[2].text_input("Gender", key="kyc_gender", placeholder="select/text")

        row2 = st.columns(3, gap="large")
        row2[0].text_input("ID Expiry Status", key="kyc_id_expired", placeholder="select")
        row2[1].text_input("ID Expiry Date", key="kyc_id_expiry", placeholder="date")
        row2[2].text_input(
            "ID Type",
            key="kyc_id_type",
            placeholder="passport / ID Card / Resident Identity Card",
        )

        row3 = st.columns(3, gap="large")
        row3[0].text_input("ID Number", key="kyc_id_number")
        row3[1].text_input("Date of Birth", key="kyc_dob", placeholder="date")
        row3[2].text_input("Submit Time", key="kyc_submit_time", placeholder="datetime")

        row4 = st.columns(3, gap="large")
        row4[0].text_input("Review Time", key="kyc_review_time", placeholder="datetime")
        row4[1].text_input("Submit IP", key="kyc_submit_ip")
        row4[2].text_input("IP Location", key="kyc_ip_location")

        row5 = st.columns(2, gap="large")
        row5[0].text_input("Submit Device ID", key="kyc_device_id")
        row5[1].text_input("Device Type", key="kyc_device_type")

        btn_spacer, btn_col = st.columns([5, 1])
        with btn_col:
            generate_clicked = st.button("Generate KYC Document", use_container_width=True, key="kyc_generate")
        if generate_clicked:
            data = {}
            for key, _label in KYC_FIELDS:
                value = st.session_state.get(f"kyc_{key}", "")
                if key == "id_type":
                    value = normalize_kyc_value(key, value)
                data[key] = value
            account_id = st.session_state.get("kyc_account_id", "")
            template_path = os.path.join(BASE_DIR, KYC_TEMPLATE_ORDERED)
            if not os.path.exists(template_path):
                template_path = os.path.join(BASE_DIR, KYC_TEMPLATE_FILE)
            safe_name = sanitize_filename_component(data.get("name", "KYC")) or "KYC"
            filename = f"{safe_name} - KYC.docx"
            images = st.session_state.get("kyc_files") or []
            doc_bytes = fill_kyc_document(template_path, data, account_id, images)
            st.session_state["kyc_download"] = {"filename": filename, "bytes": doc_bytes}
            st.success("已生成文档，请下载。")

        download_payload = st.session_state.get("kyc_download")
        if download_payload:
            st.download_button(
                "下载 KYC 文档",
                data=download_payload["bytes"],
                file_name=download_payload["filename"],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="kyc_download_btn",
            )
